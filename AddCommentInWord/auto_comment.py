# -*- coding: utf-8 -*-
"""
自动批量给 Word/PDF 作业添加评语
配置区在文件顶部，按需修改后运行即可。
"""

import os
import re
import sys
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import red
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ================= 配置区域 =================
# 务必确认路径真实存在！

# 1. Excel 评语文件完整路径
EXCEL_PATH = r"E:\AddCommentInWord\examples\comments.xlsx"

# 2. 原始作业(Word/PDF)所在文件夹
INPUT_FOLDER = r"E:\AddCommentInWord\examples"

# 3. 处理结果保存位置
OUTPUT_FOLDER = r"E:\AddCommentInWord\processed"

# 4. Windows 字体路径（PDF 写中文必备）
FONT_PATH = r"C:\Windows\Fonts\simhei.ttf"
# ==========================================


class AutoComment:
    """批量给 Word/PDF 文件添加评语"""

    def __init__(self):
        self.excel_path   = Path(EXCEL_PATH)
        self.input_folder = Path(INPUT_FOLDER)
        self.output_folder= Path(OUTPUT_FOLDER)
        self.output_folder.mkdir(parents=True, exist_ok=True)

        self.comments_dict = {}

    # ---------- 评语读取 ----------
    def load_comments(self) -> bool:
        """读取 Excel 评语表"""
        print(f"📖 读取评语表: {self.excel_path}")
        if not self.excel_path.exists():
            print(f"❌ 找不到 Excel: {self.excel_path}")
            return False

        try:
            df = pd.read_excel(self.excel_path, dtype=str)
            for _, row in df.iterrows():
                sid, comment = map(str, (row.iloc[0], row.iloc[1]))
                if sid and comment and sid.lower() != 'nan':
                    self.comments_dict[sid.strip()] = comment.strip()
            print(f"✅ 已加载 {len(self.comments_dict)} 条评语")
            return True
        except Exception as e:
            print(f"❌ 读取 Excel 失败: {e}")
            return False

    # ---------- 学号提取 ----------
    @staticmethod
    def get_student_id(filename: str):
        match = re.search(r'(\d+)', filename)
        return match.group(1) if match else None

    # ---------- Word 处理 ----------
    def process_word(self, file_path: Path, save_path: Path, comment: str) -> bool:
        """在页眉右侧写入评语"""
        try:
            doc = Document(file_path)
            header = doc.sections[0].header
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = para.add_run(f"【教师评语】\n{comment}")
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)

            doc.save(save_path)
            return True
        except Exception as e:
            print(f"   ❌ Word 处理出错: {e}")
            return False

    # ---------- PDF 处理 ----------
    def process_pdf(self, file_path: Path, save_path: Path, comment: str) -> bool:
        """在每页右上角叠加评语水印"""
        temp_watermark = self.output_folder / "temp_watermark.pdf"
        try:
            # 1. 生成水印
            c = canvas.Canvas(str(temp_watermark), pagesize=letter)

            # 注册中文字体
            if os.path.exists(FONT_PATH):
                try:
                    pdfmetrics.registerFont(TTFont('SimHei', FONT_PATH))
                    c.setFont('SimHei', 12)
                except Exception as e:
                    print(f"   ⚠️ 字体加载失败: {e}")
                    c.setFont("Helvetica", 12)
            else:
                print("   ⚠️ 未找到中文字体，可能出现乱码")
                c.setFont("Helvetica", 12)

            c.setFillColor(red)
            c.drawString(300, 750, f"Comment: {comment}")
            c.save()

            # 2. 合并到每一页
            reader = PdfReader(file_path)
            writer = PdfWriter()
            wm_page = PdfReader(str(temp_watermark)).pages[0]

            for page in reader.pages:
                page.merge_page(wm_page)
                page.compress_content_streams()
                writer.add_page(page)

            # 3. 保存
            with open(save_path, 'wb') as f:
                writer.write(f)
            return True

        except Exception as e:
            import traceback
            print(f"   ❌ PDF 处理错误: {e}")
            traceback.print_exc()
            return False
        finally:
            if temp_watermark.exists():
                try:
                    os.remove(temp_watermark)
                except: pass

    # ---------- 主流程 ----------
    def run(self):
        if not self.load_comments():
            return

        print("\n🚀 开始处理作业文件...")
        files = list(self.input_folder.glob('*.docx')) + list(self.input_folder.glob('*.pdf'))

        count = 0
        for file_path in files:
            if file_path.name.startswith('~$'):
                continue

            sid = self.get_student_id(file_path.name)
            if sid not in self.comments_dict:
                continue  # 跳过无评语匹配的文件

            print(f"正在处理: {file_path.name}")
            save_path = self.output_folder / f"{file_path.stem}_已批改{file_path.suffix}"

            success = False
            if file_path.suffix.lower() == '.docx':
                success = self.process_word(file_path, save_path, self.comments_dict[sid])
            elif file_path.suffix.lower() == '.pdf':
                success = self.process_pdf(file_path, save_path, self.comments_dict[sid])

            if success:
                print("  ✅ 完成")
                count += 1

        print(f"\n🎉 全部结束！共成功处理 {count} 个文件。")
        print(f"📂 结果保存在: {self.output_folder}")


if __name__ == "__main__":
    AutoComment().run()